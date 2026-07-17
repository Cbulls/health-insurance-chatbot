"""DOCX / HWP5 / DOC convert / 레지스트리 테스트."""
from __future__ import annotations

import io
from unittest.mock import MagicMock

import pytest

from harag.parsing.convert_service import ConvertError, ConvertService
from harag.parsing.document_parser import DocumentParser
from harag.parsing.docx_decoder import DocxDecoder
from harag.parsing.hwp5_decoder import Hwp5Decoder
from harag.parsing.parser import DecodeError


def _minimal_docx_bytes() -> bytes:
    from docx import Document
    doc = Document()
    doc.add_heading("제12조 출장비", level=1)
    doc.add_paragraph("국내 출장 한도는 1일 10만원이다.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "항목"
    table.cell(0, 1).text = "한도"
    table.cell(1, 0).text = "출장"
    table.cell(1, 1).text = "10만원"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_docx_decoder_paragraphs_and_table():
    blocks = DocxDecoder().decode(_minimal_docx_bytes(), "docx")
    kinds = {b.kind for b in blocks}
    assert "heading" in kinds or "paragraph" in kinds
    assert "table" in kinds
    texts = " ".join(b.text for b in blocks if b.text)
    assert "출장" in texts or any(
        "출장" in (c.get("text") or "") for b in blocks if b.kind == "table"
        for c in b.cells)


def test_document_parser_docx_registry():
    p = DocumentParser()
    ir = p.parse(
        _minimal_docx_bytes(),
        document_id="d-docx",
        filename="규정.docx",
    )
    assert ir.blocks
    assert ir.source_format.value == "docx"
    assert ir.parse_status.value in ("ok", "partial")


def test_hwp5_decoder_rejects_bad_bytes():
    dec = Hwp5Decoder()
    with pytest.raises(DecodeError):
        dec.decode(b"not-an-ole", "hwp5")


def test_hwp5_ole_prvtext_via_monkeypatch(monkeypatch):
    """PrvText UTF-16 추출 경로."""
    import olefile

    class FakeOle:
        def __init__(self, *a, **k):
            pass

        def exists(self, name):
            return name in ("FileHeader", "PrvText")

        def openstream(self, name):
            class S:
                def read(self_inner):
                    return "제1조 여비\n국내 출장 한도 10만원".encode("utf-16-le")
            return S()

        def listdir(self):
            return []

        def close(self):
            pass

    monkeypatch.setattr(olefile, "OleFileIO", FakeOle)
    blocks = Hwp5Decoder().decode(b"x", "hwp5")
    assert any("여비" in b.text or "출장" in b.text for b in blocks)
    assert all(b.kind in ("heading", "paragraph") for b in blocks)


def test_document_parser_hwp5_sets_table_warning(monkeypatch):
    import olefile

    class FakeOle:
        def __init__(self, *a, **k):
            pass

        def exists(self, name):
            return name in ("FileHeader", "PrvText")

        def openstream(self, name):
            class S:
                def read(self_inner):
                    return "제3조 내용\n본문입니다.".encode("utf-16-le")
            return S()

        def listdir(self):
            return []

        def close(self):
            pass

    monkeypatch.setattr(olefile, "OleFileIO", FakeOle)
    p = DocumentParser()
    ir = p.parse(b"x", document_id="h1", filename="a.hwp")
    assert ir.blocks
    assert p.last_warning == "hwp5_table_limited"
    assert ir.parse_status.value == "partial"


def test_doc_convert_service_disabled():
    svc = ConvertService(enabled=False)
    with pytest.raises(ConvertError):
        svc.doc_to_docx(b"x")


def test_document_parser_doc_uses_convert(monkeypatch):
    conv = MagicMock()
    conv.doc_to_docx.return_value = _minimal_docx_bytes()
    p = DocumentParser(convert=conv)
    ir = p.parse(b"fake-doc", document_id="d1", filename="old.doc")
    conv.doc_to_docx.assert_called_once()
    assert ir.blocks
    assert ir.extraction_path.value == "pdf_via"


def test_suffix_allowlist():
    from harag.api.routes_ingest import _suffix_of, _ALLOWED_SUFFIXES
    assert _suffix_of("a.docx") == ".docx"
    assert _suffix_of("a.doc") == ".doc"
    assert _suffix_of("a.hwp") == ".hwp"
    assert ".docx" in _ALLOWED_SUFFIXES
    assert ".doc" in _ALLOWED_SUFFIXES
